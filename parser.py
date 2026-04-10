from dotenv import load_dotenv
from google import genai
from google.genai import types
from prompts.parser_prompt import PARSER_SYSTEM_PROMPT, PARSER_USER_TEMPLATE
import json
import fitz
import docx
import os
import io


load_dotenv()

# The new client automatically picks up GOOGLE_API_KEY from the environment
client = genai.Client()

def llm_extract_fields(resume_text: str):
    user_prompt = PARSER_USER_TEMPLATE.format(resume_text=resume_text)

    response = client.models.generate_content(
        model="gemini-1.5-flash", 
        contents=user_prompt,
        config=types.GenerateContentConfig(
            system_instruction=PARSER_SYSTEM_PROMPT,
            response_mime_type="application/json",
        )
    )
    return json.loads(response.text)

def extract_text_from_pdf(file_input) -> str:
    text = ""
    if isinstance(file_input, bytes):
        with fitz.open(stream=file_input, filetype="pdf") as pdf:
            for page in pdf:
                text += page.get_text()
    else:
        with fitz.open(file_input) as pdf:
            for page in pdf:
                text += page.get_text()
    return text

def extract_text_from_docx(file_input) -> str:
    if isinstance(file_input, bytes):
        doc = docx.Document(io.BytesIO(file_input))
    else:
        doc = docx.Document(file_input)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

if __name__ == "__main__":
    # Replace this with the path to a real PDF or DOCX file
    sample_file = "Chiklit-Gohil-Resume.pdf"
    
    if os.path.exists(sample_file):
        print(f"Extracting text from {sample_file}...")
        extracted_text = extract_text(sample_file)
        print("Parsing extracted text with Gemini...\n")
        print(json.dumps(llm_extract_fields(extracted_text), indent=2))
    else:
        print(f"Please place a '{sample_file}' in the project root to test.")
